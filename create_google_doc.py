#!/usr/bin/env python3
"""
Script to create a formatted DOCX file from the markdown interview prep.
This can be directly uploaded to Google Drive and opened as a Google Doc.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

# Create a new Document
doc = Document()

# Set document properties
doc.core_properties.author = "Lee Roy Pinto"
doc.core_properties.title = "Anthropic Interview Preparation Guide"
doc.core_properties.subject = "Enterprise Account Executive, Industries Generalist - Sydney"

# Add title
title = doc.add_heading('Anthropic Interview Preparation Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_heading('Enterprise Account Executive, Industries Generalist - Sydney', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add candidate info
doc.add_paragraph()
info = doc.add_paragraph()
info.add_run('Candidate: ').bold = True
info.add_run('Lee Roy Pinto\n')
info.add_run('Location: ').bold = True
info.add_run('Sydney, Australia\n')
info.add_run('Target Start: ').bold = True
info.add_run('Q1 2025')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Read the markdown content
with open('anthropic_interview_prep.md', 'r') as f:
    content = f.read()

# Split content into sections
sections = content.split('\n## ')

# Process each section
for section in sections[1:]:  # Skip the first title section
    lines = section.split('\n')
    
    # Add section heading
    if lines[0].startswith('Part'):
        doc.add_heading(lines[0], 1)
    else:
        doc.add_heading(lines[0], 2)
    
    # Process the rest of the lines
    current_paragraph = None
    in_list = False
    list_level = 0
    
    for line in lines[1:]:
        # Skip empty lines
        if not line.strip():
            if current_paragraph:
                current_paragraph = None
            continue
        
        # Handle subsection headings (###)
        if line.startswith('### '):
            heading_text = line.replace('### ', '')
            doc.add_heading(heading_text, 3)
            current_paragraph = None
            in_list = False
            
        # Handle bold headers with colons
        elif line.startswith('**') and ':' in line:
            p = doc.add_paragraph()
            # Parse bold text and regular text
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # This is bold text
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
            current_paragraph = p
            
        # Handle bullet points
        elif line.strip().startswith('- '):
            # Determine list level by counting leading spaces
            lead_spaces = len(line) - len(line.lstrip())
            list_level = lead_spaces // 2
            
            # Clean the bullet point text
            bullet_text = line.strip()[2:]  # Remove the '- '
            
            # Create bullet point
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            if list_level > 0:
                p.paragraph_format.left_indent = Inches(0.5 * list_level)
            
            # Handle bold text within bullet points
            if '**' in bullet_text:
                p.clear()
                parts = bullet_text.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # This is bold text
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
            
            current_paragraph = None
            in_list = True
            
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            list_text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(list_text, style='List Number')
            current_paragraph = None
            
        # Handle checkboxes
        elif line.strip().startswith('- [ ]'):
            checkbox_text = line.strip()[5:]  # Remove '- [ ] '
            p = doc.add_paragraph('☐ ' + checkbox_text)
            current_paragraph = None
            
        # Handle code blocks (triple backticks)
        elif line.strip().startswith('```'):
            # Skip code block markers
            continue
            
        # Handle quoted text (starting with >)
        elif line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            p = doc.add_paragraph(quote_text)
            p.paragraph_format.left_indent = Inches(0.5)
            # Make it italic
            for run in p.runs:
                run.italic = True
            current_paragraph = None
            
        # Handle STAR examples specially
        elif line.strip().startswith('**S:') or line.strip().startswith('**T:') or \
             line.strip().startswith('**A:') or line.strip().startswith('**R:'):
            p = doc.add_paragraph()
            # Parse the STAR format
            if '**' in line:
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # This is bold text
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(line)
            current_paragraph = None
            
        # Regular paragraph text
        else:
            if current_paragraph:
                current_paragraph.add_run(' ' + line.strip())
            else:
                p = doc.add_paragraph(line.strip())
                current_paragraph = p

# Add special formatting for key sections
for paragraph in doc.paragraphs:
    # Format "Your angle:" or "Your Response:" in blue
    if 'Your angle:' in paragraph.text or 'Your Response:' in paragraph.text:
        for run in paragraph.runs:
            if 'Your angle:' in run.text or 'Your Response:' in run.text:
                run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Format section headers
    if paragraph.text.startswith('Part '):
        for run in paragraph.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(44, 62, 80)

# Save the document
doc.save('anthropic_interview_prep.docx')

print("✅ DOCX file created successfully: anthropic_interview_prep.docx")
print("\n📤 To upload to Google Docs:")
print("1. Go to https://drive.google.com")
print("2. Click 'New' → 'File upload'")
print("3. Select 'anthropic_interview_prep.docx'")
print("4. Once uploaded, double-click to open in Google Docs")
print("\n✨ The document will retain all formatting when opened in Google Docs!")